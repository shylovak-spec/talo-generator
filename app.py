import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import datetime
import os
import re
import subprocess
import requests
import tempfile
from decimal import Decimal, ROUND_HALF_UP

# Намагаємось імпортувати бібліотеку для суми прописом
try:
    from num2words import num2words
except ImportError:
    num2words = None

# ШЛЯХ ДО ШАБЛОНІВ
TPL_DIR = "" 

# ==============================================================================
# 1. ТЕХНІЧНІ ФУНКЦІЇ
# ==============================================================================

def precise_round(number):
    return float(Decimal(str(number)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))

def format_num(n):
    return f"{precise_round(n):,.2f}".replace(",", " ").replace(".", ",")

def calculate_row(price_from_st, qty, is_fop):
    # У таблиці ціна завжди чиста, податки рахуємо в підсумку
    p_unit = precise_round(price_from_st)
    row_sum = precise_round(p_unit * qty)
    return p_unit, row_sum

def amount_to_text_uk(amount):
    val = precise_round(amount)
    grn = int(val)
    kop = int(round((val - grn) * 100))
    if num2words is None:
        return f"{format_num(val)} грн."
    try:
        words = num2words(grn, lang='uk').capitalize()
        return f"{words} гривень, {kop:02d} коп."
    except:
        return f"{format_num(val)} грн."

@st.cache_data(ttl=3600)
def load_full_database_from_gsheets():
    try:
        if "gcp_service_account" not in st.secrets: return {}
        creds = Credentials.from_service_account_info(
            st.secrets["gcp_service_account"], 
            scopes=["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        )
        gc = gspread.authorize(creds)
        sh = gc.open("База_Товарів")
        full_base = {}
        for sheet in sh.worksheets():
            category_name = sheet.title
            data = sheet.get_all_records()
            items_in_cat = {}
            for row in data:
                name = str(row.get('Назва', '')).strip()
                price_raw = str(row.get('Ціна', '0')).replace(" ", "").replace(",", ".")
                try:
                    price = float(price_raw) if (price_raw and price_raw.strip() != "") else 0.0
                except:
                    price = 0.0
                if name: items_in_cat[name] = price
            if items_in_cat: full_base[category_name] = items_in_cat
        return full_base
    except Exception as e:
        st.sidebar.error(f"⚠️ Помилка бази: {e}")
        return {}

VENDORS = {
    "ТОВ «ТАЛО»": {"full": "ТОВ «ТАЛО»", "short": "Олексій КРАМАРЕНКО", "inn": "32670939", "adr": "03113, м. Київ, проспект Перемоги, будинок 68/1 офіс 62", "iban": "UA_________________________", "bank": "АТ «УКРСИББАНК»", "tax_label": "ПДВ (20%)", "tax_rate": 0.20},
    "ФОП Крамаренко Олексій Сергійович": {"full": "ФОП Крамаренко Олексій Сергійович", "short": "Олексій КРАМАРЕНКО", "inn": "3048920896", "adr": "02156 м. Київ, вул. Кіото 9, кв. 40", "iban": "UA423348510000000026009261015", "bank": "АТ «ПУМБ»", "tax_label": "6%", "tax_rate": 0.06},
    "ФОП Шилова Ксенія Вікторівна": {"full": "ФОП Шилова Ксенія Вікторівна", "short": "Ксенія ШИЛОВА", "inn": "3237308989", "adr": "20901 м. Чигирин, вул. Миру 4, кв. 43", "iban": "UA433220010000026007350102344", "bank": "АТ УНІВЕРСАЛ БАНК", "tax_label": "6%", "tax_rate": 0.06}
}

# ==============================================================================
# 2. PDF ТА ТЕЛЕГРАМ
# ==============================================================================

def docx_to_pdf_libreoffice(docx_bytes):
    with tempfile.TemporaryDirectory() as tmp_dir:
        input_path = os.path.join(tmp_dir, "temp.docx")
        with open(input_path, "wb") as f: f.write(docx_bytes)
        try:
            subprocess.run(['lowriter', '--headless', '--convert-to', 'pdf', '--outdir', tmp_dir, input_path], check=True)
            pdf_path = os.path.join(tmp_dir, "temp.pdf")
            with open(pdf_path, "rb") as f: return f.read()
        except: return None

def send_telegram_file(file_bytes, file_name):
    token = st.secrets.get("telegram_bot_token")
    chat_id = st.secrets.get("telegram_chat_id")
    if not token or not chat_id: return
    url = f"https://api.telegram.org/bot{token}/sendDocument"
    try:
        files = {'document': (file_name, file_bytes)}
        requests.post(url, files=files, data={'chat_id': chat_id})
        st.toast(f"✅ Відправлено КП в Telegram")
    except: pass

# ==============================================================================
# 3. ФОРМАТУВАННЯ ТА ЗАМІНА
# ==============================================================================

def apply_font_style(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:ascii'), 'Times New Roman')
    r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hAnsi'), 'Times New Roman')

def set_cell_style(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    run = p.add_run(str(text))
    apply_font_style(run, 12, bold, italic)

def replace_with_formatting(doc, reps):
    for p in doc.paragraphs:
        for k, v in reps.items():
            placeholder = f"{{{{{k}}}}}"
            if placeholder in p.text:
                full_text = p.text.replace(placeholder, str(v))
                p.text = ""
                if ":" in full_text:
                    parts = full_text.split(":", 1)
                    r1 = p.add_run(parts[0] + ":")
                    apply_font_style(r1, 12, bold=True)
                    r2 = p.add_run(parts[1])
                    apply_font_style(r2, 12, bold=False)
                else:
                    r = p.add_run(full_text)
                    apply_font_style(r, 12)

    for tbl in doc.tables:
        if any("Найменування" in cell.text for row in tbl.rows for cell in row.cells):
            continue 
        for row in tbl.rows:
            for cell in row.cells:
                for k, v in reps.items():
                    placeholder = f"{{{{{k}}}}}"
                    if placeholder in cell.text:
                        txt = cell.text.replace(placeholder, str(v))
                        cell.text = ""
                        p = cell.paragraphs[0]
                        if ":" in txt:
                            parts = txt.split(":", 1)
                            r1 = p.add_run(parts[0] + ":")
                            apply_font_style(r1, 12, bold=True)
                            r2 = p.add_run(parts[1])
                            apply_font_style(r2, 12, bold=False)
                        else:
                            r = p.add_run(txt)
                            apply_font_style(r, 12)

def fill_document_table(doc, items, is_fop, label_name):
    target_table = None
    for tbl in doc.tables:
        if any("Найменування" in cell.text for cell in tbl.rows[0].cells):
            target_table = tbl
            break
    if not target_table: return 0

    total_no_tax = 0
    cols = len(target_table.columns)
    is_spec = "Специфікація" in label_name

    categories = {}
    for it in items:
        cat = it['cat'].upper()
        if cat not in categories: categories[cat] = []
        categories[cat].append(it)

    for cat_name, cat_items in categories.items():
        row_cat = target_table.add_row()
        row_cat.cells[0].merge(row_cat.cells[cols-1])
        set_cell_style(row_cat.cells[0], cat_name, WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        
        for it in cat_items:
            p_unit, row_sum = calculate_row(it['p'], it['qty'], is_fop)
            total_no_tax += row_sum
            r = target_table.add_row()
            set_cell_style(r.cells[0], it['name'])
            if cols >= 4:
                set_cell_style(r.cells[1], str(it['qty']), WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_style(r.cells[2], format_num(p_unit), WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_style(r.cells[3], format_num(row_sum), WD_ALIGN_PARAGRAPH.RIGHT)

    # ЛОГІКА ПІДСУМКІВ
    tax_rate = 0.06 if is_fop else 0.20
    tax_amount = precise_round(total_no_tax * tax_rate)
    grand_total = precise_round(total_no_tax + tax_amount)

    if is_fop and is_spec:
        # Для специфікації ФОП тільки один рядок
        r = target_table.add_row()
        r.cells[0].merge(r.cells[cols-2])
        set_cell_style(r.cells[0], "ЗАГАЛЬНА СУМА, грн:", WD_ALIGN_PARAGRAPH.LEFT, True)
        set_cell_style(r.cells[cols-1], format_num(grand_total), WD_ALIGN_PARAGRAPH.RIGHT, True)
    else:
        # Для ТОВ (КП/Спец) та для ФОП (КП) - три рядки
        sub_label = "РАЗОМ (без навантаження), грн:" if is_fop else "РАЗОМ (без ПДВ), грн:"
        tax_label = "Податкове навантаження 6%:" if is_fop else "ПДВ (20%):"
        total_label = "ЗАГАЛЬНА СУМА, грн:" if is_fop else "ЗАГАЛЬНА СУМА з ПДВ, грн:"
        
        f_rows = [(sub_label, total_no_tax, False), (tax_label, tax_amount, False), (total_label, grand_total, True)]
        for lab, val, bld in f_rows:
            r = target_table.add_row()
            r.cells[0].merge(r.cells[cols-2])
            set_cell_style(r.cells[0], lab, WD_ALIGN_PARAGRAPH.LEFT, bld)
            set_cell_style(r.cells[cols-1], format_num(val), WD_ALIGN_PARAGRAPH.RIGHT, bld)

    return grand_total

# ==============================================================================
# 4. ІНТЕРФЕЙС STREAMLIT
# ==============================================================================

st.set_page_config(page_title="Talo Generator", layout="wide")
st.title("⚡ Генератор КП")

EQUIPMENT_BASE = load_full_database_from_gsheets()
if "generated_files" not in st.session_state: st.session_state.generated_files = None

with st.expander("📌 Основні дані", expanded=True):
    c1, c2 = st.columns(2)
    vendor_choice = c1.selectbox("Виконавець:", list(VENDORS.keys()))
    is_fop = "ФОП" in vendor_choice
    v = VENDORS[vendor_choice]
    customer = c1.text_input("Замовник", "ОСББ")
    address = c1.text_input("Адреса об'єкта", "м. Київ")
    kp_num = c2.text_input("Номер КП", "1223.25")
    manager = c2.text_input("Відповідальний", "Олексій Крамаренко")
    date_str = c2.date_input("Дата", datetime.date.today()).strftime("%d.%m.%Y")
    phone = c2.text_input("Телефон", "+380 (67) 477-17-18")
    email = c2.text_input("E-mail", "o.kramarenko@talo.com.ua")

st.subheader("📝 Текст для КП")
txt_intro = st.text_area("Вступний текст", "Відповідно до наданих даних пропонуємо наступне:")
tc1, tc2, tc3 = st.columns(3)
l1 = tc1.text_input("Пункт 1", "Організація автономного живлення ліфтів")
l2 = tc2.text_input("Пункт 2", "Організація автономного живлення насосної")
l3 = tc3.text_input("Пункт 3", "Аварійне освітлення та відеонагляд")

st.subheader("📦 Специфікація")
items_to_generate = []
if EQUIPMENT_BASE:
    tabs = st.tabs(list(EQUIPMENT_BASE.keys()))
    for i, cat in enumerate(EQUIPMENT_BASE.keys()):
        with tabs[i]:
            sel = st.multiselect(f"Додати з {cat}:", list(EQUIPMENT_BASE[cat].keys()), key=f"ms_{cat}")
            for name in sel:
                base_p = EQUIPMENT_BASE[cat][name]
                cn, cq, cp = st.columns([4, 1, 2])
                cn.write(f"**{name}**")
                q = cq.number_input("К-сть", 1, 500, 1, key=f"qty_{cat}_{name}")
                p = cp.number_input("Ціна (змінюйте)", 0.0, 1000000.0, float(base_p), key=f"prc_{cat}_{name}")
                items_to_generate.append({"name": name, "qty": q, "p": p, "cat": cat})

if items_to_generate:
    # ПРАВКА 2: ВИВЕДЕННЯ СУМИ НА СТРІМЛІТ
    total_raw = sum(it['p'] * it['qty'] for it in items_to_generate)
    tax_rate = 0.06 if is_fop else 0.20
    tax_val = precise_round(total_raw * tax_rate)
    final_sum = total_raw + tax_val
    
    st.markdown(f"""
    <div style="background-color:#f0f2f6; padding:20px; border-radius:10px; border-left: 5px solid #ff4b4b;">
        <h3 style="margin-top:0;">💰 Попередній розрахунок:</h3>
        <p>Сума без податку: <b>{format_num(total_raw)} грн.</b></p>
        <p>Податок ({'6%' if is_fop else '20%'}): <b>{format_num(tax_val)} грн.</b></p>
        <h2 style="color:#ff4b4b;">РАЗОМ: {format_num(final_sum)} грн.</h2>
    </div>
    """, unsafe_allow_html=True)

    st.write("---")
    c_gen, c_tg = st.columns(2)
    
    if c_gen.button("📄 1. ЗГЕНЕРУВАТИ ДОКУМЕНТИ", use_container_width=True):
        reps = {"vendor_name": v["full"], "vendor_address": v["adr"], "vendor_inn": v["inn"], "vendor_iban": v["iban"], 
                "vendor_bank": v["bank"], "vendor_email": email, "vendor_short_name": v["short"], "customer": customer, 
                "address": address, "kp_num": kp_num, "date": date_str, "manager": manager, "phone": phone, "email": email,
                "txt_intro": txt_intro, "line1": l1, "line2": l2, "line3": l3, "spec_id_postavka": kp_num, "spec_id_roboti": kp_num}
        
        try:
            creds = Credentials.from_service_account_info(st.secrets["gcp_service_account"], scopes=["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"])
            gspread.authorize(creds).open("Реєстр КП Talo").get_worksheet(0).append_row([date_str, kp_num, customer, address, vendor_choice, final_sum, manager])
        except: pass

        results = {}
        file_map = {"КП": f"{TPL_DIR}template.docx", "Специфікація_ОБЛ": f"{TPL_DIR}template_postavka.docx", "Специфікація_РОБ": f"{TPL_DIR}template_roboti.docx"}
        clean_addr = re.sub(r'[^\w\s-]', '', address).replace(' ', '_')[:30]

        for label, full_tpl_path in file_map.items():
            if os.path.exists(full_tpl_path):
                doc = Document(full_tpl_path)
                it_fill = items_to_generate
                if "ОБЛ" in label: it_fill = [i for i in items_to_generate if "роботи" not in i["cat"].lower()]
                if "РОБ" in label: it_fill = [i for i in items_to_generate if "роботи" in i["cat"].lower()]
                
                if it_fill:
                    actual_total = fill_document_table(doc, it_fill, is_fop, label)
                    reps["total_sum_digits"] = format_num(actual_total)
                    reps["total_sum_words"] = amount_to_text_uk(actual_total)
                    replace_with_formatting(doc, reps)
                    buf = BytesIO(); doc.save(buf); buf.seek(0)
                    results[label] = {"name": f"{label}_{kp_num}_{clean_addr}.docx", "data": buf}
        
        st.session_state.generated_files = results
        st.rerun()

    if c_tg.button("✈️ 2. ВІДПРАВИТИ КП В TELEGRAM (PDF)", use_container_width=True, type="primary"):
        if st.session_state.generated_files and "КП" in st.session_state.generated_files:
            info = st.session_state.generated_files["КП"]
            pdf_data = docx_to_pdf_libreoffice(info['data'].getvalue())
            if pdf_data: send_telegram_file(pdf_data, info['name'].replace(".docx", ".pdf"))
        else:
            st.warning("Спочатку згенеруйте документи")

if st.session_state.generated_files:
    cols = st.columns(len(st.session_state.generated_files))
    for i, (k, info) in enumerate(st.session_state.generated_files.items()):
        cols[i].download_button(f"💾 {info['name']}", info['data'], info['name'])
